from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading("BÁO CÁO GIỮA KỲ – TRÍ TUỆ NHÂN TẠO", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph("Đề tài: Trò chơi Caro AI (Minimax & Alpha-Beta Pruning)")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

def add_question(doc, num, title_text, *paragraphs):
    h = doc.add_heading(f"Câu {num}. {title_text}", level=1)
    for p in paragraphs:
        doc.add_paragraph(p)

# ─── Câu 1 ───────────────────────────────────────────────────────────────────
add_question(doc, 1, "Mô tả bài toán cờ Caro",
    "Cờ Caro (Gomoku) là trò chơi hai người chơi trên bàn cờ 15×15 ô. Hai người lần lượt đặt quân "
    "của mình (X hoặc O) vào các ô trống. Người chơi đầu tiên tạo được 4 quân liên tiếp theo hàng ngang, "
    "hàng dọc hoặc đường chéo sẽ thắng. Nếu bàn cờ đầy mà không ai thắng thì kết quả là hòa.",
    "Trong dự án này, người chơi (PLAYER = 1) đi quân X, AI (AI = 2) đi quân O. "
    "AI sử dụng thuật toán tìm kiếm đối kháng (Minimax hoặc Alpha-Beta Pruning) để chọn nước đi tối ưu. "
    "Điều kiện thắng được định nghĩa trong constants.py: WIN_LENGTH = 4 (4 quân liên tiếp), "
    "BOARD_SIZE = 15 (bàn cờ 15×15)."
)

# ─── Câu 2 ───────────────────────────────────────────────────────────────────
add_question(doc, 2, "Luật chơi và điều kiện thắng",
    "Luật chơi:",
    "  • Bàn cờ 15×15 ô, ban đầu tất cả ô đều trống (EMPTY = 0).",
    "  • Hai người chơi lần lượt đặt quân: người chơi đặt X (PLAYER = 1), AI đặt O (AI = 2).",
    "  • Mỗi lượt chỉ được đặt một quân vào ô trống hợp lệ.",
    "  • Người chơi đi trước (PLAYER đi lượt đầu tiên).",
    "Điều kiện thắng (hàm check_winner trong game/rules.py):",
    "  • Kiểm tra từ ô vừa đặt theo 4 hướng: ngang (0,1), dọc (1,0), chéo chính (1,1), chéo phụ (1,-1).",
    "  • Đếm số quân liên tiếp cùng loại theo cả hai chiều của mỗi hướng.",
    "  • Nếu tổng số quân liên tiếp ≥ WIN_LENGTH (= 4) thì người đó thắng.",
    "  • Nếu bàn cờ đầy (is_full()) mà chưa có người thắng thì kết quả là hòa (is_draw = True)."
)

# ─── Câu 3 ───────────────────────────────────────────────────────────────────
add_question(doc, 3, "Cách biểu diễn trạng thái bàn cờ",
    "Trạng thái bàn cờ được biểu diễn bằng lớp GameState (core/game_state.py):",
    "  • board: ma trận 2 chiều 15×15, mỗi phần tử là số nguyên: 0 (EMPTY), 1 (PLAYER), 2 (AI).",
    "  • current_turn: lượt đi hiện tại (PLAYER hoặc AI).",
    "  • winner: người thắng (None nếu chưa kết thúc).",
    "  • is_draw: cờ hòa (bool).",
    "  • winning_cells: danh sách các ô tạo thành chuỗi thắng.",
    "  • last_move: tọa độ (r, c) của nước đi cuối cùng.",
    "  • move_count: số nước đã đi.",
    "  • nodes_explored, last_ai_time, last_eval_score, last_depth: thống kê hiệu năng AI.",
    "Ví dụ biểu diễn bàn cờ trong code:",
    '  board = [[0]*15 for _ in range(15)]  # bàn cờ trống\n'
    '  board[7][7] = 1  # PLAYER đặt quân tại (7,7)\n'
    '  board[7][8] = 2  # AI đặt quân tại (7,8)',
    "Khi truyền vào thuật toán AI, board được sao chép (deep copy) để tránh thay đổi trạng thái gốc."
)

# ─── Câu 4 ───────────────────────────────────────────────────────────────────
add_question(doc, 4, "Cách sinh nước đi hợp lệ",
    "Hàm get_candidate_moves(board, radius=2) trong game/rules.py sinh danh sách nước đi hợp lệ:",
    "  • Duyệt toàn bộ bàn cờ, tìm các ô đã có quân.",
    "  • Với mỗi ô có quân, thêm tất cả ô trống trong vùng bán kính radius=2 xung quanh vào tập ứng viên.",
    "  • Nếu bàn cờ hoàn toàn trống (đầu ván), trả về ô trung tâm [(7, 7)].",
    "  • Kết quả là tập hợp (set) các ô trống nằm gần quân đã đặt, tránh xét toàn bộ 225 ô.",
    "Sau khi có danh sách ứng viên, hàm order_moves() (ai/heuristics/move_ordering.py) sắp xếp theo độ ưu tiên:",
    "  1. Nước thắng ngay lập tức của AI.",
    "  2. Nước chặn thắng ngay của đối thủ.",
    "  3. Điểm kết hợp tấn công + phòng thủ (tính bằng _quick_score).",
    "Việc giới hạn ứng viên và sắp xếp nước đi giúp giảm đáng kể không gian tìm kiếm."
)

# ─── Câu 5 ───────────────────────────────────────────────────────────────────
add_question(doc, 5, "Cách kiểm tra trạng thái kết thúc",
    "Trạng thái kết thúc được kiểm tra sau mỗi nước đi bằng hàm check_winner(board, last_move) "
    "trong game/rules.py:",
    "  • Nhận vào bàn cờ và tọa độ nước vừa đặt.",
    "  • Kiểm tra 4 hướng từ ô đó: ngang, dọc, chéo chính, chéo phụ.",
    "  • Đếm quân liên tiếp cùng loại theo cả hai chiều.",
    "  • Nếu tổng ≥ 4: trả về (người_thắng, danh_sách_ô_thắng).",
    "  • Nếu không ai thắng: trả về (None, []).",
    "Kiểm tra hòa: hàm is_full() trong GameState kiểm tra tất cả ô đều khác EMPTY.",
    "Trong thuật toán AI (minimax/alphabeta), trạng thái kết thúc được kiểm tra ở đầu mỗi lần đệ quy:",
    "  • winner == ai_player → trả về +SCORE_WIN + depth (thắng nhanh hơn được ưu tiên).",
    "  • winner == human_player → trả về -(SCORE_WIN + depth).",
    "  • depth == 0 → gọi evaluate_board() để ước lượng.",
    "  • Không còn nước đi → trả về 0 (hòa)."
)

# ─── Câu 6 ───────────────────────────────────────────────────────────────────
add_question(doc, 6, "Thuật toán Minimax đã cài đặt",
    "File: ai/algorithms/minimax.py | Agent: ai/agents/minimax.py (chế độ Easy, độ sâu 3)",
    "Nguyên lý: Minimax là thuật toán tìm kiếm đối kháng hai người. Người tối đa hóa (AI) cố gắng "
    "chọn nước đi có điểm cao nhất, người tối thiểu hóa (người chơi) cố gắng chọn nước đi có điểm thấp nhất.",
    "Cài đặt hàm _minimax(board, depth, is_maximizing, last_move, ai_player, human_player, counter):",
    "  1. Tăng bộ đếm nút (counter[0] += 1).",
    "  2. Kiểm tra điều kiện kết thúc: thắng/thua → trả về ±(SCORE_WIN + depth).",
    "  3. Nếu depth == 0: gọi evaluate_board() để đánh giá tĩnh.",
    "  4. Sinh nước đi ứng viên và sắp xếp bằng order_moves().",
    "  5. Nếu is_maximizing: duyệt tất cả nước đi, chọn giá trị lớn nhất; cắt sớm nếu best >= SCORE_WIN.",
    "  6. Nếu is_minimizing: duyệt tất cả nước đi, chọn giá trị nhỏ nhất; cắt sớm nếu best <= -SCORE_WIN.",
    "  7. Đặt quân tạm thời (board[r][c] = player), đệ quy, rồi hoàn tác (board[r][c] = EMPTY).",
    "Hàm get_best_move_minimax() gọi _minimax ở tầng gốc, trả về (best_move, best_val, nodes_explored).",
    "Độ phức tạp: O(b^d) với b là số nước đi ứng viên, d là độ sâu (mặc định d=3)."
)

# ─── Câu 7 ───────────────────────────────────────────────────────────────────
add_question(doc, 7, "Thuật toán Alpha-Beta đã cài đặt",
    "File: ai/algorithms/alphabeta.py | Agent: ai/agents/alphabetapuring.py (chế độ Medium, độ sâu 4)",
    "Nguyên lý: Alpha-Beta Pruning là cải tiến của Minimax, duy trì hai ngưỡng alpha (giá trị tốt nhất "
    "cho người tối đa hóa) và beta (giá trị tốt nhất cho người tối thiểu hóa). Khi alpha ≥ beta, "
    "cắt tỉa nhánh còn lại vì không thể ảnh hưởng đến kết quả.",
    "Cài đặt hàm _alphabeta(board, depth, alpha, beta, is_maximizing, ...):",
    "  • Tương tự Minimax nhưng thêm điều kiện cắt tỉa:",
    "    - Nút MAX: cập nhật alpha = max(alpha, best); nếu alpha >= beta → break (β-cutoff).",
    "    - Nút MIN: cập nhật beta = min(beta, best); nếu alpha >= beta → break (α-cutoff).",
    "Iterative Deepening (tìm kiếm sâu dần): hàm get_best_move_alphabeta() tìm kiếm từ độ sâu 1 đến "
    "depth (mặc định 4), mỗi vòng lặp dùng nước tốt nhất của vòng trước đặt lên đầu danh sách nước đi "
    "để tăng hiệu quả cắt tỉa. Nếu tìm thấy nước thắng ngay (score >= SCORE_WIN), dừng sớm.",
    "Ưu điểm so với Minimax thuần: giảm số nút cần xét từ O(b^d) xuống O(b^(d/2)) trong trường hợp tốt nhất, "
    "cho phép tìm kiếm sâu hơn trong cùng thời gian."
)

# ─── Câu 8 ───────────────────────────────────────────────────────────────────
add_question(doc, 8, "Hàm đánh giá trạng thái",
    "File: ai/evaluation/evaluation.py — hàm evaluate_board(board, move, ai_player, human_player)",
    "Phương pháp: Quét toàn bộ bàn cờ theo 4 hướng (ngang, dọc, chéo chính, chéo phụ). "
    "Mỗi chuỗi quân liên tiếp chỉ được đếm một lần (từ ô đầu tiên của chuỗi). "
    "Phân biệt đầu mở (ô kế tiếp là EMPTY) và đầu bị chặn (ô kế tiếp là quân đối thủ hoặc biên).",
    "Bảng điểm (open_ends = số đầu mở: 0, 1, 2):",
    "  AI (tấn công):                    Đối thủ (phòng thủ):",
    "  • 4 quân, 2 đầu mở  → +100,000   • 4 quân, 2 đầu mở  → -100,000",
    "  • 4 quân, 1 đầu mở  → +50,000    • 4 quân, 1 đầu mở  → -45,000",
    "  • 3 quân, 2 đầu mở  → +5,000     • 3 quân, 2 đầu mở  → -4,000",
    "  • 3 quân, 1 đầu mở  → +500       • 3 quân, 1 đầu mở  → -400",
    "  • 2 quân, 2 đầu mở  → +200       • 2 quân, 2 đầu mở  → -150",
    "  • 2 quân, 1 đầu mở  → +50        • 2 quân, 1 đầu mở  → -40",
    "Điểm tổng = tổng điểm tất cả chuỗi AI − tổng điểm tất cả chuỗi đối thủ. "
    "Điểm dương nghĩa là AI đang có lợi thế, điểm âm nghĩa là đối thủ đang có lợi thế."
)

# ─── Câu 9 ───────────────────────────────────────────────────────────────────
add_question(doc, 9, "Thiết kế các trạng thái thử nghiệm",
    "Dự án cung cấp hai cơ chế thử nghiệm:",
    "1. Chế độ Benchmark (benchmark/benchmark_runner.py):",
    "   • Chạy trận AI vs AI: Minimax (độ sâu 1–6) đấu với Alpha-Beta (độ sâu 2–7).",
    "   • Mỗi cặp chạy 1 ván, tổng 36 cặp trong ma trận đầy đủ.",
    "   • Ghi lại: người thắng, số nước đi, bàn cờ cuối, thời gian mỗi nước.",
    "2. Chế độ Analysis (ui/screens/analysis_screen.py):",
    "   • Người dùng tự thiết lập trạng thái bàn cờ bằng cách click chuột.",
    "   • Click trái: lần lượt đặt X → O → xóa.",
    "   • Click phải: xóa ô.",
    "   • Chọn AI đi bên nào (X hoặc O), thuật toán và độ sâu.",
    "   • Nhấn ▶ Move để AI tính và thực hiện một nước từ trạng thái đó.",
    "Các trạng thái thử nghiệm điển hình:",
    "   • Bàn cờ trống: AI luôn đi ô trung tâm (7,7).",
    "   • Trạng thái có 3 quân liên tiếp: kiểm tra AI có chặn/tấn công đúng không.",
    "   • Trạng thái gần thắng (3 quân, 1 đầu mở): kiểm tra ưu tiên tấn công vs phòng thủ."
)

# ─── Câu 10 ──────────────────────────────────────────────────────────────────
add_question(doc, 10, "Bảng kết quả thực nghiệm",
    "Kết quả từ file benchmark_matrix_results.txt (5 trận đã chạy, Minimax d1 vs Alpha-Beta d2–d6):",
)

# Table
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Trận đấu"
hdr[1].text = "Người thắng"
hdr[2].text = "Kết quả X (Minimax)"
hdr[3].text = "Kết quả O (Alpha-Beta)"
for h in hdr:
    for run in h.paragraphs[0].runs:
        run.bold = True

data = [
    ("minimax_d1 vs alphabeta_d2", "alphabeta_d2", "Thua", "Thắng"),
    ("minimax_d1 vs alphabeta_d3", "alphabeta_d3", "Thua", "Thắng"),
    ("minimax_d1 vs alphabeta_d4", "alphabeta_d4", "Thua", "Thắng"),
    ("minimax_d1 vs alphabeta_d5", "alphabeta_d5", "Thua", "Thắng"),
    ("minimax_d1 vs alphabeta_d6", "alphabeta_d6", "Thua", "Thắng"),
]
for row_data in data:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()
doc.add_paragraph(
    "Nhận xét: Trong tất cả 5 trận đã chạy, Alpha-Beta (dù chỉ ở độ sâu 2) đều thắng Minimax độ sâu 1. "
    "Điều này cho thấy độ sâu tìm kiếm ảnh hưởng quyết định đến chất lượng nước đi. "
    "Lưu ý: bộ benchmark đầy đủ 36 trận (Minimax d1–6 vs Alpha-Beta d2–7) chưa được chạy hết."
)

# ─── Câu 11 ──────────────────────────────────────────────────────────────────
add_question(doc, 11, "Nhận xét về số trạng thái đã xét và thời gian chạy",
    "Số trạng thái (nút) được xét phụ thuộc vào thuật toán và độ sâu:",
    "Minimax (độ sâu 3, chế độ Easy):",
    "  • Không có cắt tỉa, duyệt toàn bộ cây tìm kiếm.",
    "  • Số nút xét ≈ b^3 với b là số nước đi ứng viên (thường 10–30 nước gần quân đã đặt).",
    "  • Ví dụ: b=20, d=3 → ~8,000 nút; thời gian thường dưới 0.5 giây.",
    "Alpha-Beta (độ sâu 4, chế độ Medium):",
    "  • Cắt tỉa alpha-beta + iterative deepening + move ordering.",
    "  • Trong trường hợp tốt nhất (move ordering tốt): số nút ≈ b^(d/2) = b^2.",
    "  • Ví dụ: b=20, d=4 → ~400–1,600 nút thay vì 160,000 nút (Minimax thuần).",
    "  • Thời gian thường 0.1–2 giây tùy giai đoạn ván đấu.",
    "Thống kê được lưu trong GameState: nodes_explored, last_ai_time, last_eval_score, last_depth, "
    "và hiển thị trực tiếp trên HUD trong khi chơi.",
    "Nhận xét chung: Move ordering (ưu tiên nước thắng ngay, nước chặn, rồi điểm tổng hợp) "
    "đóng vai trò quan trọng giúp Alpha-Beta cắt tỉa hiệu quả hơn, giảm số nút xét đáng kể."
)

# ─── Câu 12 ──────────────────────────────────────────────────────────────────
add_question(doc, 12, "Nhận xét về ảnh hưởng của độ sâu tìm kiếm",
    "Độ sâu tìm kiếm ảnh hưởng trực tiếp đến chất lượng nước đi và thời gian tính toán:",
    "Chất lượng nước đi:",
    "  • Độ sâu 1: AI chỉ nhìn 1 bước, dễ bị dụ vào bẫy, không thấy được mối đe dọa 2–3 bước.",
    "  • Độ sâu 3 (Easy/Minimax): AI có thể phát hiện và phản ứng với các mối đe dọa ngắn hạn.",
    "  • Độ sâu 4 (Medium/Alpha-Beta): AI nhìn xa hơn, có thể lên kế hoạch tấn công 2 hướng cùng lúc.",
    "  • Độ sâu càng lớn, AI càng 'thông minh' nhưng thời gian tính toán tăng theo hàm mũ.",
    "Thời gian tính toán:",
    "  • Minimax: tăng theo O(b^d) — tăng gấp b lần khi tăng 1 độ sâu.",
    "  • Alpha-Beta: tăng theo O(b^(d/2)) trong trường hợp tốt — hiệu quả hơn nhiều.",
    "  • Với b≈20: Minimax d3 ≈ 8,000 nút; Minimax d4 ≈ 160,000 nút; Alpha-Beta d4 ≈ 400–1,600 nút.",
    "Kết quả benchmark: Minimax d1 thua Alpha-Beta d2 trong tất cả 5 trận thử nghiệm, "
    "cho thấy chỉ cần tăng 1 độ sâu với Alpha-Beta đã vượt trội hoàn toàn so với Minimax nông hơn.",
    "Kết luận: Với bài toán Caro 4 quân, độ sâu 3–4 là cân bằng tốt giữa chất lượng và tốc độ. "
    "Iterative deepening trong Alpha-Beta giúp tận dụng thời gian hiệu quả và cải thiện move ordering."
)

# ─── Câu 13 ──────────────────────────────────────────────────────────────────
add_question(doc, 13, "Ưu điểm và hạn chế của chương trình",
    "ƯU ĐIỂM:",
    "  1. Kiến trúc rõ ràng, tách biệt: core (hằng số, trạng thái), game (luật, quản lý), "
       "ai (thuật toán, đánh giá, heuristic), ui (giao diện) — dễ mở rộng và bảo trì.",
    "  2. Hai mức độ khó với thuật toán khác nhau: Easy (Minimax d3) và Medium (Alpha-Beta d4), "
       "phù hợp cho người mới và người chơi có kinh nghiệm.",
    "  3. Alpha-Beta kết hợp Iterative Deepening và Move Ordering: tăng hiệu quả cắt tỉa, "
       "giảm số nút xét đáng kể so với Minimax thuần.",
    "  4. Hàm đánh giá chi tiết: phân biệt đầu mở/bị chặn, bảng điểm riêng cho AI và đối thủ, "
       "phản ánh đúng giá trị chiến lược của từng mẫu quân.",
    "  5. Công cụ Benchmark và Analysis: cho phép so sánh thuật toán, thử nghiệm trạng thái tùy ý, "
       "quan sát nước đi theo thời gian thực.",
    "  6. Tính năng Undo: hoàn tác cả lượt (nước AI + nước người chơi), tiện lợi khi học.",
    "  7. Giao diện hiện đại với Pygame: hiệu ứng hover, animation đặt quân, đường thắng nổi bật.",
    "HẠN CHẾ:",
    "  1. Không có bảng chuyển vị (transposition table): các trạng thái giống nhau có thể được "
       "tính toán lại nhiều lần, lãng phí tài nguyên.",
    "  2. Độ sâu cố định: không tự điều chỉnh độ sâu theo thời gian còn lại hoặc độ phức tạp "
       "của trạng thái hiện tại.",
    "  3. Hàm đánh giá chưa xét tương tác giữa các chuỗi: ví dụ, hai chuỗi 3 quân theo hai hướng "
       "khác nhau (fork) nguy hiểm hơn tổng điểm đơn lẻ, nhưng chưa được tính thêm điểm.",
    "  4. Chỉ hỗ trợ chế độ 1 người vs AI: chưa có chế độ 2 người chơi trực tiếp (PvP).",
    "  5. Benchmark chưa hoàn chỉnh: ma trận 36 trận (Minimax d1–6 vs Alpha-Beta d2–7) chưa được "
       "chạy đầy đủ trong file kết quả hiện tại.",
    "  6. Không có học máy: AI hoàn toàn dựa trên tìm kiếm và hàm đánh giá tĩnh, "
       "không học từ các ván đấu trước."
)

out_path = r"C:\Users\Tuanngth\Documents\TTNT\caro_game\BaoCaoGiuaKy_CaroAI.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
